"""
Парсер DOCX файлов для извлечения текста
"""
import io
from typing import Optional
from docx import Document
from docx.opc.exceptions import PackageNotFoundError


class DocxParser:
    """
    Парсер для извлечения текста из DOCX файлов
    """
    
    @staticmethod
    def extract_text(file_content: bytes) -> Optional[str]:
        """
        Извлекает текст из DOCX файла
        
        Args:
            file_content: Содержимое DOCX файла в байтах
            
        Returns:
            Извлеченный текст или None при ошибке
        """
        try:
            # Создаем поток из байтов
            file_stream = io.BytesIO(file_content)
            
            # Загружаем документ
            doc = Document(file_stream)
            
            # Извлекаем текст из всех параграфов
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            # Объединяем все части текста
            full_text = "\n\n".join(text_parts)
            
            return full_text if full_text.strip() else None
            
        except PackageNotFoundError:
            # Файл не является валидным DOCX
            return None
        except Exception:
            # Любая другая ошибка при парсинге
            return None
    
    @staticmethod
    def validate_docx(file_content: bytes) -> bool:
        """
        Проверяет, является ли файл валидным DOCX
        
        Args:
            file_content: Содержимое файла в байтах
            
        Returns:
            True если файл валидный DOCX, False иначе
        """
        try:
            file_stream = io.BytesIO(file_content)
            Document(file_stream)
            return True
        except (PackageNotFoundError, Exception):
            return False
    
    @staticmethod
    def get_document_info(file_content: bytes) -> dict:
        """
        Получает информацию о документе
        
        Args:
            file_content: Содержимое DOCX файла в байтах
            
        Returns:
            Словарь с информацией о документе
        """
        try:
            file_stream = io.BytesIO(file_content)
            doc = Document(file_stream)
            
            # Подсчитываем статистику
            paragraphs_count = len([p for p in doc.paragraphs if p.text.strip()])
            tables_count = len(doc.tables)
            
            # Пытаемся получить метаданные
            core_props = doc.core_properties
            
            return {
                "paragraphs_count": paragraphs_count,
                "tables_count": tables_count,
                "title": core_props.title or "",
                "author": core_props.author or "",
                "created": core_props.created,
                "modified": core_props.modified,
            }
            
        except Exception:
            return {
                "paragraphs_count": 0,
                "tables_count": 0,
                "title": "",
                "author": "",
                "created": None,
                "modified": None,
            }
